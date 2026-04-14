"""
Triplet — Document Parser
Extracts raw text from PDF, DOCX, and image-based PDFs (OCR)
"""
import io

import pdfplumber
import pytesseract
from PIL import Image
from docx import Document
from pypdf import PdfReader


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """
    Extract text from a PDF file.
    Falls back to OCR if text layer is empty (scanned PDF).
    """
    text = ""
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except Exception:
        pass

    # Secondary fallback for PDFs that pdfplumber struggles with.
    if not text.strip():
        try:
            reader = PdfReader(io.BytesIO(file_bytes))
            for page in reader.pages:
                page_text = page.extract_text() or ""
                if page_text:
                    text += page_text + "\n"
        except Exception:
            pass

    # If no text extracted, try OCR
    if not text.strip():
        text = _ocr_pdf(file_bytes)

    return text.strip()


def _ocr_pdf(file_bytes: bytes) -> str:
    """
    OCR fallback for image-based PDFs.
    Converts each page to image and runs tesseract.
    """
    text = ""
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        for page in doc:
            pix = page.get_pixmap(dpi=200)
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            page_text = pytesseract.image_to_string(img)
            text += page_text + "\n"
    except Exception:
        pass
    return text.strip()


def extract_text_from_image(file_bytes: bytes) -> str:
    """
    Extract text directly from JPG/PNG uploads using OCR.
    """
    try:
        with Image.open(io.BytesIO(file_bytes)) as img:
            normalized = img.convert("RGB")
            return pytesseract.image_to_string(normalized).strip()
    except Exception:
        return ""


def extract_text_from_docx(file_bytes: bytes) -> str:
    """
    Extract text from a DOCX file.
    """
    text = ""
    try:
        doc = Document(io.BytesIO(file_bytes))
        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text + "\n"
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + "\n"
    except Exception:
        pass
    return text.strip()


def extract_text(file_bytes: bytes, mime_type: str) -> str:
    """
    Main entry point. Routes to correct extractor based on mime type.
    Returns clean plain text.
    """
    if mime_type == "application/pdf":
        return extract_text_from_pdf(file_bytes)
    elif mime_type == (
        "application/vnd.openxmlformats-officedocument"
        ".wordprocessingml.document"
    ):
        return extract_text_from_docx(file_bytes)
    elif mime_type in {"image/jpeg", "image/png"}:
        return extract_text_from_image(file_bytes)
    else:
        # Try PDF as default
        return extract_text_from_pdf(file_bytes)
